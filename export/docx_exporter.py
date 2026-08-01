from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generation.image_budget import count_body_chinese_chars
from generation.content_quality import contamination_hits
from export.layout_pipeline import prepare_article_layout

DEFAULT_FONT = "宋体"
TITLE_FONT = "黑体"
SUBTITLE_FONT = "楷体"
SOURCE_FONT = "宋体"
ARTICLE_NOT_READY_MESSAGE = "文章尚未生成成功，暂时无法导出。请先重试文章生成。"


def _set_run_font(run: Any, font_name: str, size: int, *, bold: bool = False, color: str | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_font(paragraph: Any, font_name: str, size: int, *, bold: bool = False, color: str | None = None) -> None:
    for run in paragraph.runs:
        _set_run_font(run, font_name, size, bold=bold, color=color)


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _resolve(path: str | None, base_dir: Path | None) -> Path | None:
    if not path:
        return None
    candidate = Path(path)
    if not candidate.is_absolute() and base_dir is not None:
        candidate = base_dir / candidate
    return candidate


def _add_image(paragraph: Any, path: str | None, base_dir: Path | None, width: float) -> bool:
    resolved = _resolve(path, base_dir)
    if not resolved or not resolved.is_file():
        return False
    paragraph.add_run().add_picture(str(resolved), width=Inches(width))
    return True


def _image_index(article: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {
        str(item.get("paragraph_ref")): item
        for item in article.get("images") or []
        if item.get("role") == "inline" and item.get("paragraph_ref")
    }


def _body_markdown_from_sections(article: dict[str, Any]) -> str:
    if str(article.get("body_markdown") or "").strip():
        return str(article.get("body_markdown") or "").strip()
    parts: list[str] = []
    for section in article.get("sections") or []:
        if not isinstance(section, dict):
            continue
        body = str(section.get("body") or "").strip()
        if body:
            parts.append(body)
    return "\n\n".join(parts).strip()


def ensure_article_ready_for_docx_export(article: dict[str, Any] | None) -> None:
    if not isinstance(article, dict):
        raise ValueError(f"ARTICLE_NOT_READY: {ARTICLE_NOT_READY_MESSAGE}")
    title = str(article.get("title") or "").strip()
    body_markdown = _body_markdown_from_sections(article)
    body_char_count = int(article.get("body_char_count") or count_body_chinese_chars(article) or 0)
    quality_gate = article.get("quality_gate") if isinstance(article.get("quality_gate"), dict) else {}
    if (
        not title
        or not body_markdown
        or body_char_count <= 0
        or article.get("used_local_fallback")
        or article.get("fallback_kind")
        or article.get("article_quality_blocked")
        or str(article.get("recommended_status") or "") in {"retry_required", "quality_blocked"}
        or str(quality_gate.get("status") or "") == "failed"
        or contamination_hits(str(article.get("content_markdown") or body_markdown))
    ):
        raise ValueError(f"ARTICLE_NOT_READY: {ARTICLE_NOT_READY_MESSAGE}")


def _format_body_paragraph(paragraph: Any, *, first_line: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line:
        fmt.first_line_indent = Pt(24)
    _set_paragraph_font(paragraph, DEFAULT_FONT, 12)


def _configure(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    title = document.styles["Title"]
    title.font.name = TITLE_FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), TITLE_FONT)
    title.font.size = Pt(20)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = SUBTITLE_FONT
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), SUBTITLE_FONT)
    subtitle.font.size = Pt(12)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_after = Pt(10)

    heading = document.styles["Heading 1"]
    heading.font.name = TITLE_FONT
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), TITLE_FONT)
    heading.font.size = Pt(15)
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)


def _caption(document: Document, text: str) -> None:
    if not text:
        return
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    _set_paragraph_font(paragraph, DEFAULT_FONT, 9, color="666666")


def _add_source_paragraph(document: Document, source: str) -> None:
    lines = [line.strip() for line in str(source or "").splitlines() if line.strip()]
    if not lines:
        return
    first = document.add_paragraph(lines[0])
    first.paragraph_format.space_after = Pt(2)
    first.paragraph_format.first_line_indent = Pt(0)
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_font(first, SOURCE_FONT, 10)
    for line in lines[1:]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if line.startswith("原文链接："):
            label, _, url = line.partition("：")
            run = paragraph.add_run(f"{label}：")
            _set_run_font(run, SOURCE_FONT, 10)
            _add_hyperlink(paragraph, url.strip(), url.strip())
        else:
            run = paragraph.add_run(line)
            _set_run_font(run, SOURCE_FONT, 10)


def _add_article_content(document: Document, article: dict[str, Any], base_dir: Path | None) -> None:
    article = prepare_article_layout(article)
    heading = document.add_paragraph(article["title"], style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(heading, TITLE_FONT, 20, bold=True)

    subtitle = str(article.get("subtitle") or article.get("intro") or "").strip()
    if subtitle:
        paragraph = document.add_paragraph(subtitle, style="Subtitle")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(10)
        _set_paragraph_font(paragraph, SUBTITLE_FONT, 12)

    cover = next((item for item in article.get("images") or [] if item.get("role") == "cover" and item.get("status") == "completed"), None)
    if cover:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _add_image(paragraph, cover.get("path"), base_dir, 6.2):
            _caption(document, str(cover.get("caption") or "封面图片"))

    inline_images = _image_index(article)
    for index, section in enumerate(article.get("sections") or [], start=1):
        title = document.add_paragraph(str(section.get("heading") or f"正文 {index}"), style="Heading 1")
        _set_paragraph_font(title, TITLE_FONT, 15, bold=True)
        paragraphs = [block.strip() for block in re.split(r"\n\s*\n+", str(section.get("body") or "")) if block.strip()]
        for paragraph_text in paragraphs:
            paragraph = document.add_paragraph(paragraph_text)
            _format_body_paragraph(paragraph)
        image = inline_images.get(f"section-{index}")
        if image and image.get("status") == "completed":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if _add_image(paragraph, image.get("path"), base_dir, 5.8):
                _caption(document, str(image.get("caption") or f"配图 {index}"))

    keywords = article.get("keywords") or []
    if keywords:
        paragraph = document.add_paragraph("关键词：" + "、".join(str(item) for item in keywords))
        paragraph.paragraph_format.first_line_indent = Pt(0)
        _set_paragraph_font(paragraph, DEFAULT_FONT, 11)

    sources = [str(item).strip() for item in (article.get("source_list") or []) if str(item).strip()][:5]
    if sources:
        title = document.add_paragraph("资料来源", style="Heading 1")
        _set_paragraph_font(title, TITLE_FONT, 15, bold=True)
        for source in sources:
            _add_source_paragraph(document, source)



def export_article(article: dict[str, Any], output_path: Path, base_dir: Path | None = None) -> Path:
    ensure_article_ready_for_docx_export(article)
    document = Document()
    _configure(document)
    _add_article_content(document, article, base_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def export_combined(articles: list[dict[str, Any]], output_path: Path, base_dir: Path | None = None) -> Path:
    for article in articles:
        ensure_article_ready_for_docx_export(article)
    document = Document()
    _configure(document)
    for index, article in enumerate(articles):
        if index:
            document.add_page_break()
        _add_article_content(document, article, base_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
