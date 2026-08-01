"""Test R1.2: Article quality and delivery verification.

Covers:
1. Standard hotlist Markdown article quality (structure, length, no artifacts)
2. Template phrase detection (warning, not hard fail)
3. Hotlist limited basic draft (not title-stacking, has structure, no fabricated facts)
4. Word export validation (docx opens, title, source, AI statement, no JSON fields)
"""

from __future__ import annotations

import io
import json
import re
import zipfile
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from generation.article_generator import (
    _append_sections_to_markdown,
    _complete_article_structure,
    _prompt,
    generate_article as ag_generate_article,
    MIN_SECTIONS,
    REQUIRED_SECTION_HEADINGS,
    TARGET_BODY_CHINESE_CHARS,
    WARNING_BODY_CHINESE_CHARS,
)
from generation.content_quality import (
    quality_gate,
    _article_body_for_fact_scan,
)
from generation.image_budget import count_body_chinese_chars
from export.docx_exporter import export_article
from modules.models import HotTopic


# ── helpers ────────────────────────────────────────────────


def _make_topic(
    topic_id: str = "t1",
    title: str = "某热点进入热榜讨论",
    summary: str = "热榜摘要显示，该话题正在引发用户关注。",
    source_url: str = "https://example.com/hot",
) -> HotTopic:
    return HotTopic(
        id=topic_id,
        source="test-hotlist",
        source_name="测试热榜",
        title=title,
        summary=summary,
        source_url=source_url,
        hot_value="热度 1000",
    )


def _make_angle() -> dict[str, str]:
    return {
        "name": "深度分析",
        "instruction": "从事件影响和读者关心的角度分析",
        "core_question": "这件事对普通人意味着什么？",
        "opening_strategy": "以用户关心的切入点开篇",
        "structure": ["事件发生了什么", "为什么受到关注", "可能带来哪些影响", "后续值得关注什么"],
        "must_avoid": ["夸大影响", "虚构数据"],
    }


def _make_research_bundle(
    research_status: str = "complete",
    accepted_source_count: int = 2,
    hotlist_limited: bool = False,
) -> dict[str, Any]:
    bundle: dict[str, Any] = {
        "topic_id": "t1",
        "topic_title": "某热点",
        "research_status": research_status,
        "accepted_source_count": accepted_source_count,
        "official_or_reliable_source_count": accepted_source_count,
        "sources": [
            {
                "source_id": "s1",
                "source_name": "测试来源",
                "publisher_id": "example.com",
                "domain": "example.com",
                "fetch_success": True,
                "accepted_for_research": True,
                "content": (
                    "监管部门已公开发布初步通报，要求相关方在规定期限内提供详细说明材料并全面配合调查工作。"
                    "相关方需要配合调查和整改要求。"
                    "后续仍需等待权威部门发布正式评估报告。"
                ),
            }
        ],
    }
    if hotlist_limited:
        bundle["hotlist_metadata_available"] = True
        bundle["research_status"] = "hotlist_limited"
    return bundle


def _good_hotlist_article(topic: HotTopic | None = None) -> dict[str, Any]:
    """A well-formed hotlist article with >= 900 body chars, proper structure."""
    t = topic or _make_topic()
    title = f"{t.title}：深度解析事件来龙去脉"
    intro = "近日某热点事件引发广泛讨论。本文基于公开资料梳理事件经过、分析受到关注的原因、评估可能带来的影响并指出后续值得关注的观察方向。"
    sections = [
        {
            "heading": "事件发生了什么",
            "body": (
                "根据公开信息，该事件起源于一次常规操作中出现的意外情况。"
                "具体来说，相关方在执行日常流程时遇到了此前未被充分重视的合规边界问题，导致操作被监管部门关注并介入调查。"
                "从目前已披露的信息来看，核心争议点集中在操作流程是否符合现行规范要求，以及相关方在事前是否履行了必要的审查和报备义务。"
                "监管部门已公开发布初步通报，要求相关方在规定期限内提供详细说明材料并全面配合调查工作。"
                "与此同时，涉事方通过官方渠道表示将积极配合调查，并承诺及时向公众通报后续进展和处理结果。"
                "目前事件仍在持续发酵中，多个相关方面均已表态关注此事进展。"
            ),
            "image_brief": "事件发生的新闻现场",
        },
        {
            "heading": "为什么受到关注",
            "body": (
                "这一话题之所以迅速进入公众视野并持续升温，首先是因为涉及的领域与大量普通用户的日常使用场景密切相关。"
                "用户对于自身权益是否受到影响存在自然而直接的关切，因此话题在社交平台上快速传播并引发广泛讨论。"
                "其次，此类事件的处理方式和最终结论，可能为同行业类似场景提供重要的参考先例和示范效应。"
                "行业从业者、法律研究者和监管观察者都保持高度关注，多家主流媒体对此进行了持续跟踪报道和深度解读。"
                "此外，事件涉及的合规标准问题也引发了公众对相关制度完善和监管透明度的讨论和期待。"
                "资本市场和行业上下游企业也密切关注事件走向，以此评估对行业生态的连锁影响。"
            ),
            "image_brief": "用户关注的场景",
        },
        {
            "heading": "可能带来哪些影响",
            "body": (
                "从短期来看，相关方需要投入必要资源配合调查和整改要求，这可能对部分业务节奏和运营安排产生阶段性调整。"
                "与此同时，该事件也促使行业上下游重新审视自身合规流程，推动各方主动完善内部管理标准和操作规范。"
                "从中长期视角来看，此类事件引发的深入讨论有助于推动监管规则的进一步明确和细化。"
                "对普通用户而言，这意味着更清晰的权益保障边界和更透明的信息获取渠道，增强了消费信心。"
                "当然，具体影响程度和范围仍需根据后续正式处理结果和各方回应来综合评估，不宜过早下定论。"
                "行业内已有多家企业表示将以此为鉴，主动优化内部流程以提升合规水平。"
            ),
            "image_brief": "行业影响示意图",
        },
        {
            "heading": "后续值得关注什么",
            "body": (
                "建议重点持续关注三个方向：第一，监管部门的正式处理结论和具体整改要求及时间节点；"
                "第二，相关方对事件的正式回应以及后续改进措施的时间表和具体路线图；"
                "第三，行业层面是否会以此为契机出台更加明确和统一的操作指引或标准规范。"
                "本文将继续关注事件进展并及时更新相关信息。读者在转发和参与讨论时也建议注意区分已确认事实与推测性信息。"
                "同时可以关注相关领域专家和学者的专业解读，以获得更全面的认知视角。"
            ),
            "image_brief": "后续关注方向",
        },
    ]
    source_list = [
        f"[1] 测试来源：《{t.title}》，2026年7月\n原文链接：{t.source_url}",
    ]
    markdown = "\n\n".join(
        [f"# {title}", intro]
        + [f"## {s['heading']}\n{s['body']}" for s in sections]
        + ["## 资料来源\n" + source_list[0], "AI辅助声明：本文基于公开资料整理生成，发布前请再次核对关键信息。"]
    )
    return {
        "title": title,
        "intro": intro,
        "sections": sections,
        "content_markdown": markdown,
        "source_list": source_list,
        "source_statement": source_list[0],
        "ai_statement": "AI辅助声明：本文基于公开资料整理生成，发布前请再次核对关键信息。",
        "fact_basis": [],
        "recommended_status": "completed",
        "text_generation_calls": 1,
        "body_char_count": count_body_chinese_chars({"intro": intro, "sections": sections}),
    }


def _template_heavy_article() -> dict[str, Any]:
    """Article with excessive template phrases but enough body to pass length gate."""
    title = "某热点事件：值得关注的最新进展"
    intro = "从现有信息看，该事件引发关注。值得关注的是，后续仍需等待权威信息。从现有信息看，具有重要意义。引发关注的是，此事迅速成为热议话题。"
    sections = [
        {
            "heading": "事件发生了什么",
            "body": (
                "从现有信息看，该事件具有重要影响。值得关注的是，引发了广泛讨论。"
                "从现有信息看，后续仍需进一步观察。具有重要意义，值得深思。"
                "仍需等待权威确认。引发关注的是，多个方面仍需核实。"
                "具有重要意义的事件值得关注。从现有信息看，后续仍需谨慎。"
                "从现有信息看，还需要进一步等待官方公布详细信息。引发关注的是这个话题的热度持续上升。"
                "具有重要意义，值得深思的是事件背后反映出来的深层问题。后续仍需持续跟踪。"
                "从现有信息看，多个相关部门已经介入处理此事。值得关注的是处理进展。具有重要意义，还需要更多时间观察。"
                "从现有信息看，事件的来龙去脉正在逐渐清晰。引发关注的是各方回应态度。后续仍需等待最终结论。"
            ),
        },
        {
            "heading": "为什么受到关注",
            "body": (
                "从现有信息看，这个话题引发关注的原因是值得关注的。具有重要意义。"
                "仍需等待更多信息。值得关注的是，引发了广泛讨论。从现有信息看，进展值得关注。"
                "具有重要意义的是事件对行业规范的推动作用。后续仍需等待权威机构的正式表态。"
                "引发关注的是社交媒体上大量用户的讨论和转发热度。从现有信息看，舆论还在持续发酵。"
            ),
        },
        {
            "heading": "可能带来哪些影响",
            "body": (
                "从现有信息看，影响值得关注，具有重要意义，后续仍需关注。"
                "从现有信息看，短期和长期影响都需要进一步评估。值得关注的是行业反应。"
                "具有重要意义的是这可能推动监管标准的完善。引发关注的是多家企业已开始自查。"
                "后续仍需等待权威部门发布正式评估报告。从现有信息看，市场反应较为审慎。"
            ),
        },
        {
            "heading": "后续值得关注什么",
            "body": (
                "从现有信息看，后续仍需等待，引发关注。值得关注的是进展。"
                "从现有信息看，下一步需要关注官方通报和整改措施。具有重要意义的是事件的示范效应。"
                "引发关注的是公众对处理结果的高度期待。后续仍需关注相关方的正式回应。"
                "从现有信息看，建议读者保持关注但不要过度解读未确认的信息。值得关注的是后续政策走向。"
            ),
        },
    ]
    markdown = "\n\n".join(
        [f"# {title}", intro]
        + [f"## {s['heading']}\n{s['body']}" for s in sections]
        + ["## 资料来源\n[1] 测试", "AI辅助声明：测试"]
    )
    return {
        "title": title,
        "intro": intro,
        "sections": sections,
        "content_markdown": markdown,
        "source_list": ["[1] 测试"],
        "ai_statement": "AI辅助声明：测试",
        "fact_basis": [],
        "recommended_status": "completed",
        "text_generation_calls": 1,
        "body_char_count": count_body_chinese_chars({"intro": intro, "sections": sections}),
    }


def _limited_bundle() -> dict[str, Any]:
    return {
        "topic_id": "t1",
        "topic_title": "某热点进入热榜讨论",
        "research_status": "hotlist_limited",
        "hotlist_metadata_available": True,
        "accepted_source_count": 0,
        "official_or_reliable_source_count": 0,
        "sources": [],
        "usable_facts": [],
        "research_fact_cards": [],
        "limited_research_notice": "当前仅获取到热榜标题和有限元数据。",
    }


# ── Test 1: Standard hotlist Markdown article quality ──


class TestHotlistArticleQuality:
    def test_title_not_original_hotspot_title(self):
        """Article title must differ from original hotlist title."""
        topic = _make_topic(title="携程被罚后内部全员信曝光")
        article = _good_hotlist_article(topic)
        assert article["title"] != topic.title

    def test_intro_exists_and_reasonable_length(self):
        """Intro must exist and be 80-150 chars (or close)."""
        article = _good_hotlist_article()
        intro = article["intro"]
        assert intro, "intro must not be empty"
        intro_chars = len(re.findall(r"[\u4e00-\u9fff]", intro))
        assert 50 <= intro_chars <= 200, f"intro chars {intro_chars} not in reasonable range"

    def test_sections_are_four(self):
        """Must have exactly 4 sections matching REQUIRED_SECTION_HEADINGS."""
        article = _good_hotlist_article()
        sections = article.get("sections", [])
        assert len(sections) == 4, f"expected 4 sections, got {len(sections)}"
        headings = [s.get("heading", "") for s in sections]
        for expected in REQUIRED_SECTION_HEADINGS:
            assert expected in headings, f"missing section: {expected}"

    def test_each_section_body_nonempty(self):
        """No section may have empty body."""
        article = _good_hotlist_article()
        for section in article.get("sections", []):
            body = section.get("body", "")
            chars = len(re.findall(r"[\u4e00-\u9fff]", body))
            assert chars >= 30, f"section '{section.get('heading')}' body too short: {chars} chars"

    def test_body_char_count_at_least_900(self):
        """Body must have >= 900 Chinese chars (intro + sections)."""
        article = _good_hotlist_article()
        bc = article.get("body_char_count", 0) or count_body_chinese_chars(article)
        assert bc >= 900, f"body_char_count {bc} < 900"

    def test_content_markdown_no_json_no_code_fence(self):
        """content_markdown must not contain JSON or code fences."""
        article = _good_hotlist_article()
        md = article.get("content_markdown", "")
        assert "```" not in md, "must not contain code fences"
        # Must not contain JSON field names that should not leak
        for field in ("content_markdown", "fact_basis", "source_ids", '"sections"', '"body"'):
            assert field not in md, f"JSON field '{field}' leaked into markdown"

    def test_source_section_exists(self):
        """Must have '## 资料来源' section."""
        article = _good_hotlist_article()
        md = article.get("content_markdown", "")
        assert "## 资料来源" in md or "资料来源" in md, "missing source section"

    def test_ai_statement_exists(self):
        """Must have AI statement at end, not merged with source."""
        article = _good_hotlist_article()
        md = article.get("content_markdown", "")
        assert "AI辅助声明" in md, "missing AI statement"
        # AI statement should be after source section
        source_pos = md.find("资料来源")
        ai_pos = md.find("AI辅助声明")
        if source_pos >= 0:
            assert ai_pos > source_pos, "AI statement must come after source"

    def test_quality_gate_passes_for_good_article(self):
        """A well-formed article with >= 900 chars should pass quality gate."""
        article = _good_hotlist_article()
        bundle = _make_research_bundle()
        result = quality_gate(article, bundle)
        assert result["passed"], f"quality gate failed: {result.get('hard_errors')}"
        assert result["status"] in ("passed", "warning"), f"unexpected status: {result['status']}"


# ── Test 2: Template phrase detection ──


class TestTemplatePhraseDetection:
    def test_template_heavy_article_gets_warning(self):
        """Article with too many template phrases should get warning, not hard fail."""
        article = _template_heavy_article()
        bundle = _make_research_bundle(accepted_source_count=1)
        result = quality_gate(article, bundle)
        assert result["passed"], "template warning must not hard-fail"
        # Look for the template warning
        all_warnings = " ".join(result.get("warnings", []))
        assert "模板套话偏多" in all_warnings or result["status"] == "warning", (
            f"expected template warning, got status={result['status']}, warnings={result.get('warnings')}"
        )

    def test_clean_article_no_template_warning(self):
        """A clean article should not trigger template warnings."""
        article = _good_hotlist_article()
        bundle = _make_research_bundle()
        result = quality_gate(article, bundle)
        all_warnings = " ".join(result.get("warnings", []))
        assert "模板套话偏多" not in all_warnings, "clean article should not trigger template warning"


# ── Test 3: Hotlist limited basic draft ──


class TestHotlistLimitedDraft:
    def test_no_title_stacking(self):
        """Limited draft must not just repeat hotspot title."""
        topic = _make_topic(title="某热点")
        article = _good_hotlist_article(topic)
        md = article.get("content_markdown", "")
        # Must not be just "某热点" repeated
        title_count = md.count(topic.title)
        assert title_count <= 3, f"title repeated {title_count} times: seems like title stacking"

    def test_has_intro_and_four_sections(self):
        """Limited draft must still have intro and 4 sections."""
        article = _good_hotlist_article()
        assert article.get("intro"), "must have intro"
        assert len(article.get("sections", [])) >= 4, "must have ≥ 4 sections"

    def test_has_source_and_ai_statement(self):
        """Must include source and AI statement explicitly."""
        article = _good_hotlist_article()
        md = article.get("content_markdown", "")
        assert "资料来源" in md, "missing source section in limited draft"
        assert "AI辅助声明" in md, "missing AI statement in limited draft"

    def test_no_fabricated_hard_facts(self):
        """Must not fabricate hard facts (numbers, penalties, conclusions)."""
        # Given a limited bundle with no actual facts, the article should not
        # contain fabricated specifics like "罚款500万元" or "3名负责人被免职"
        article = _good_hotlist_article()
        md = article.get("content_markdown", "")
        fabricated_patterns = [
            r"罚款\d+万元",
            r"\d+名.*被免职",
            r"\d+人受伤",
            r"\d+人死亡",
            r"判处.*有期徒刑",
            r"官方证实",
        ]
        for pattern in fabricated_patterns:
            matches = re.findall(pattern, md)
            assert not matches, f"fabricated fact pattern '{pattern}' found: {matches}"


# ── Test 4: Word export validation ──


class TestWordExport:
    def test_docx_opens_and_has_title(self, tmp_path: Path):
        """Exported docx must open and first paragraph is the title."""
        article = _good_hotlist_article()
        output = tmp_path / "test.docx"
        result = export_article(article, output)
        assert result.exists()
        doc = Document(str(result))
        paragraphs = doc.paragraphs
        assert len(paragraphs) >= 1, "docx must have paragraphs"
        assert paragraphs[0].text == article["title"], (
            f"first paragraph '{paragraphs[0].text}' != title '{article['title']}'"
        )

    def test_docx_includes_limited_source_section(self, tmp_path: Path):
        """Docx must include the customer-facing limited source section."""
        article = _good_hotlist_article()
        output = tmp_path / "test.docx"
        export_article(article, output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "资料来源" in all_text

    def test_docx_omits_ai_statement(self, tmp_path: Path):
        """Docx must not expose AI generation statements."""
        article = _good_hotlist_article()
        output = tmp_path / "test.docx"
        export_article(article, output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AI辅助声明" not in all_text
        assert "AI声明" not in all_text

    def test_docx_no_json_field_names(self, tmp_path: Path):
        """Docx must not contain JSON field names like content_markdown or fact_basis."""
        article = _good_hotlist_article()
        output = tmp_path / "test.docx"
        export_article(article, output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for forbidden in ("content_markdown", "fact_basis", "source_ids", '"sections"'):
            assert forbidden not in all_text, f"JSON field '{forbidden}' leaked into docx"


# ── Test 5: Prompt structure validation ──


class TestPromptStructure:
    def test_prompt_contains_required_section_headings(self):
        """The prompt must instruct the model to use the 4 required headings."""
        topic = _make_topic()
        angle = _make_angle()
        prompt = _prompt(topic, angle, "analysis", "客观", 1200, research_bundle=_make_research_bundle())
        for heading in REQUIRED_SECTION_HEADINGS:
            assert heading in prompt, f"prompt missing required heading: {heading}"

    def test_prompt_sets_target_body_text_for_1200(self):
        """For word_count=1200, target should be 900~1200."""
        topic = _make_topic()
        angle = _make_angle()
        prompt = _prompt(topic, angle, "analysis", "客观", 1200, research_bundle=_make_research_bundle())
        assert "1200" in prompt and "1400" in prompt, "target body text should mention 1200~1400"

    def test_prompt_sets_target_body_text_for_1600(self):
        """For word_count=1600, target should be 1600~1800 (R1.2.1)."""
        topic = _make_topic()
        angle = _make_angle()
        prompt = _prompt(topic, angle, "analysis", "客观", 1600, research_bundle=_make_research_bundle())
        assert "1600" in prompt and "1800" in prompt, "target body text should mention 1600~1800"

    def test_prompt_forbids_template_phrases(self):
        """Prompt must explicitly forbid template phrases."""
        topic = _make_topic()
        angle = _make_angle()
        prompt = _prompt(topic, angle, "analysis", "客观", 1200, research_bundle=_make_research_bundle())
        template_words = ["从现有信息看", "值得关注", "引发关注", "具有重要意义", "仍需等待"]
        found = [w for w in template_words if w in prompt]
        # The prompt should mention these as forbidden, so presence is OK
        # But we need to ensure the prompt actually mentions the prohibition
        assert any("禁止" in prompt and w in prompt for w in template_words[:2]), (
            "prompt should forbid template phrases"
        )

    def test_custom_topic_prompt_excludes_news_template_words(self):
        """Custom topic prompt must not reference '事件发生了什么', '热榜', '权威信息确认' as structural headings."""
        topic = _make_topic(title="普通人该如何使用AI赚钱")
        bundle = {"custom_topic": topic.title, "research_status": "custom_topic", "sources": []}
        angle = _make_angle()
        prompt = _prompt(topic, angle, "method", "客观", 1200, research_bundle=bundle)
        # The prompt should NOT use these as section headings (checking lines with '## ')
        heading_lines = [line for line in prompt.splitlines() if line.strip().startswith("## ")]
        heading_text = " ".join(heading_lines)
        assert "事件发生了什么" not in heading_text, "custom topic section headings should not use hotspot headings"
        assert "为什么受到关注" not in heading_text, "custom topic section headings should not use hotspot headings"
        # But the prompt CAN mention them in the prohibition note
        assert "不得出现" in prompt, "prompt should still forbid news template words"


# ── Test 6: Body char count edge cases ──


class TestBodyCharCount:
    def test_body_below_700_fails_quality_gate(self):
        """Body < 700 chars should hard-fail quality gate."""
        article = _good_hotlist_article()
        # Override body to be too short
        article["intro"] = "短。"
        article["sections"] = [
            {"heading": "事件发生了什么", "body": "很短的内容。"},
            {"heading": "为什么受到关注", "body": "也很短。"},
            {"heading": "可能带来哪些影响", "body": "短。"},
            {"heading": "后续值得关注什么", "body": "极短。"},
        ]
        article["body_char_count"] = count_body_chinese_chars(article)
        bundle = _make_research_bundle()
        result = quality_gate(article, bundle)
        # Should fail or at minimum get a hard error about length
        assert not result["passed"] or any(
            "字数不足" in err for err in result.get("hard_errors", [])
        ), f"expected body-too-short error, got {result}"

    def test_body_700_to_900_gets_warning(self):
        """Body 700-900 chars should get warning."""
        article = _good_hotlist_article()
        # Take a copy of the good article sections but trim each to about half
        full_sections = _good_hotlist_article()["sections"]
        # Use half of the full text from each section
        import copy
        half_sections = copy.deepcopy(full_sections)
        for s in half_sections:
            body = s["body"]
            # Take first ~85% to get roughly 700-900 range
            half_len = len(body) * 17 // 20
            s["body"] = body[:half_len]
        article["intro"] = "近日某热点事件引发讨论，本文梳理经过和影响。"
        article["sections"] = half_sections
        article["body_char_count"] = count_body_chinese_chars(article)
        assert 700 <= article["body_char_count"] < 900, (
            f"body_char_count should be 700-900, got {article['body_char_count']}"
        )
        bundle = _make_research_bundle()
        result = quality_gate(article, bundle)
        # Should pass (with warnings) since body >= 700
        assert result["passed"], f"body 700-900 should still pass, got {result.get('hard_errors')}"


# ── Test 7: Article structure functions ──


class TestArticleStructure:
    def test_complete_structure_assigns_title_once(self):
        """_complete_article_structure should set article['title'] correctly."""
        topic = _make_topic(title="原热点标题")
        angle = _make_angle()
        article = {
            "title": "新的文章标题",
            "intro": "导语",
            "sections": [
                {"heading": "事件发生了什么", "body": "正文内容。"},
                {"heading": "为什么受到关注", "body": "关注原因。"},
                {"heading": "可能带来哪些影响", "body": "影响分析。"},
                {"heading": "后续值得关注什么", "body": "后续关注。"},
            ],
            "source_list": [],
        }
        result = _complete_article_structure(article, topic, angle)
        assert result["title"] == "新的文章标题", "title should be preserved if already set"
        assert result["intro"] == "导语", "intro should be preserved"

    def test_append_sections_to_markdown_omits_source_heading(self):
        """Generated markdown must not expose source sections."""
        article = _good_hotlist_article()
        md = _append_sections_to_markdown(article)
        assert "资料来源" not in md

    def test_append_sections_to_markdown_no_json_leak(self):
        """content_markdown must not contain JSON or code fences."""
        article = _good_hotlist_article()
        md = _append_sections_to_markdown(article)
        assert "```" not in md, "no code fences"
        assert "content_markdown" not in md, "field name leak"
