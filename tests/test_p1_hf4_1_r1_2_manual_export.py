from __future__ import annotations

from pathlib import Path
import sys
import zipfile

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import generation.single_task as single_task
import api
import modules.generation_store as generation_store
from export.docx_exporter import export_article
from export.zip_exporter import export_article_bundle
from modules.generation_store import save_generation_task
from modules.database import SQLiteStore
from modules.models import HotTopic
from providers.text_provider import ProviderError


MANUAL_TITLE = "普通人该如何使用AI赚钱"


def _manual_topic(topic_id: str = "r1-2-manual") -> HotTopic:
    return HotTopic(
        id=topic_id,
        source="manual",
        source_name="手动输入",
        provider_status="manual",
        title=MANUAL_TITLE,
        summary="用户希望生成一篇方法型文章，说明可落地路径、步骤和风险。",
        source_url="",
    )


def _hotlist_topic(topic_id: str = "r1-2-hot") -> HotTopic:
    return HotTopic(
        id=topic_id,
        source="test-hotlist",
        source_name="测试热榜",
        title="某热点进入热榜讨论",
        summary="热榜摘要显示，该话题正在引发用户关注，更多权威信息仍待确认。",
        source_url="https://example.com/hot",
        hot_value="热度 1000",
    )


def _task(store: SQLiteStore, topic: HotTopic) -> dict:
    store.save_topics([topic])
    return store.create_task(
        f"R1.2 {topic.id}",
        "multi_topic",
        [topic.to_dict()],
        1,
        generation_options={"article_type": "科普解读", "style": "客观通俗", "image_plan_mode": "none", "word_count": 800},
    )


def _zero_bundle(topic: HotTopic) -> dict:
    return {
        "topic_id": topic.id,
        "topic_title": topic.title,
        "research_status": "not_collected",
        "accepted_source_count": 0,
        "official_or_reliable_source_count": 0,
        "candidate_link_count": 0,
        "rejected_source_count": 0,
        "sources": [],
        "usable_facts": [],
        "research_fact_cards": [],
    }


def _method_article(topic: HotTopic) -> dict:
    sections = [
        {"heading": "核心概念", "body": "普通人使用AI赚钱，关键不是追逐工具名，而是把工具能力转成别人愿意付费的服务。先明确自己能解决什么问题，再决定用AI提升效率。", "image_brief": "方法概念"},
        {"heading": "可执行方法", "body": "适合从内容整理、短文案、表格处理、客服话术、简历优化和小型自动化开始。这些需求边界清楚、交付周期短，也更容易形成案例。", "image_brief": "方法列表"},
        {"heading": "具体步骤", "body": "第一步选定一个细分场景，第二步做出样例，第三步发布服务说明，第四步接一笔小订单，第五步把交付流程沉淀成模板。", "image_brief": "步骤清单"},
        {"heading": "风险提醒", "body": "不要承诺稳定高收入，不要使用未授权素材，不要把模型结果不经检查直接交付。工具订阅、时间成本和返工风险都要提前算清楚。", "image_brief": "风险控制"},
        {"heading": "总结", "body": "先从一项简单服务跑通闭环，比同时学习很多工具更重要。有了真实反馈后，再逐步提高报价、扩展服务和自动化程度。", "image_brief": "总结"},
    ]
    source = f"[1] 手动输入：《{topic.title}》，\n原文链接："
    markdown = "\n\n".join(
        [f"# {topic.title}，真正可落地的几条路径", "导语：这篇文章面向没有技术背景的普通读者，梳理使用AI获得收入的可执行路径、交付步骤和风险边界。"]
        + [f"## {item['heading']}\n{item['body']}" for item in sections]
        + [f"## 资料来源\n{source}", "AI辅助声明：本文根据用户手动话题和AI辅助生成，发布前请补充案例、数据和来源。"]
    )
    return {
        "title": f"{topic.title}，真正可落地的几条路径",
        "intro": "导语：这篇文章面向没有技术背景的普通读者，梳理使用AI获得收入的可执行路径、交付步骤和风险边界。",
        "summary": topic.summary,
        "sections": sections,
        "content_markdown": markdown,
        "source_list": [source],
        "source_statement": source,
        "ai_statement": "AI辅助声明：本文根据用户手动话题和AI辅助生成，发布前请补充案例、数据和来源。",
        "fact_basis": [],
        "recommended_status": "completed",
        "text_generation_calls": 1,
        "text_generation_limit": 1,
    }


def _run(tmp_path: Path, monkeypatch, topic: HotTopic, generate):
    monkeypatch.setattr(generation_store, "TASKS_ROOT", tmp_path / "tasks")
    monkeypatch.setattr(single_task.ResearchService, "collect", lambda self, topic, references=None, supplemental_text="": _zero_bundle(topic))
    monkeypatch.setattr(single_task, "generate_article", generate)
    monkeypatch.setattr(single_task, "analyze_source_overlap", lambda article, bundle: {"status": "passed", "violations": []})
    store = SQLiteStore(tmp_path / f"{topic.id}.sqlite")
    task = _task(store, topic)
    return single_task.run_single_task(
        task,
        {"api_key": "saved-text-key", "has_api_key": True, "model": "r1-2-text", "timeout_seconds": 30},
        {"api_key": "image-key", "model": "test-image"},
        settings={"network": {}, "image_plan_mode": "none"},
        store=store,
    )


def test_MANUAL_TOPIC_ZERO_SOURCE_USES_TEXT_MODEL_METHOD_ARTICLE_AND_WORD(tmp_path: Path, monkeypatch):
    topic = _manual_topic()
    seen = {}

    def fake_generate(topic, angle, article_type, style, word_count, profile, **kwargs):
        seen["api_key"] = profile.get("api_key")
        seen["bundle_status"] = (kwargs.get("research_bundle") or {}).get("research_status")
        return _method_article(topic)

    result = _run(tmp_path, monkeypatch, topic, fake_generate)
    markdown = result["article"]["content_markdown"]
    assert result["status"] == "completed"
    assert result["text_generation_calls"] == 1
    assert seen == {"api_key": "saved-text-key", "bundle_status": "custom_topic"}
    assert result["research_bundle"]["research_status"] == "custom_topic"
    assert "hotlist_limited" not in str(result["research_bundle"])
    assert all(word not in markdown for word in ("热榜", "事件概览", "权威信息有限", "事件发生了什么"))
    assert "可执行方法" in markdown and "具体步骤" in markdown and "风险提醒" in markdown
    output = export_article(result["article"], tmp_path / "manual-method.docx")
    assert Document(output).paragraphs[0].text == result["article"]["title"]


def test_MANUAL_TOPIC_TIMEOUT_USES_CUSTOM_TOPIC_FALLBACK_AND_EXPORTS(tmp_path: Path, monkeypatch):
    topic = _manual_topic("r1-2-manual-timeout")
    result = _run(tmp_path, monkeypatch, topic, lambda *args, **kwargs: (_ for _ in ()).throw(ProviderError("TIMEOUT", "model timeout")))
    article = result["article"]
    markdown = article["content_markdown"]
    assert result["status"] == "completed"
    assert result["fallback_kind"] == "custom_topic_fallback"
    assert result["provider_error_code"] == "TIMEOUT"
    assert "本篇未使用文本模型正式正文" in result["fallback_notice"]
    assert all(word not in markdown for word in ("热榜", "事件概览", "等待权威信息确认", "事件发生了什么"))
    assert "核心概念" in markdown and "可执行方法" in markdown and "风险提醒" in markdown
    assert export_article(article, tmp_path / "manual-fallback.docx").is_file()


def test_HOTLIST_ZERO_SOURCE_STILL_USES_HOTLIST_LIMITED(tmp_path: Path, monkeypatch):
    topic = _hotlist_topic()
    result = _run(tmp_path, monkeypatch, topic, lambda *args, **kwargs: (_ for _ in ()).throw(ProviderError("TIMEOUT", "model timeout")))
    assert result["status"] == "completed"
    assert result["research_bundle"]["research_status"] == "hotlist_limited"
    assert result["article"]["fallback_kind"] == "hotlist_limited_draft"
    assert "事件概览" in result["article"]["content_markdown"]


def test_DESKTOP_DOWNLOADS_AND_WORD_ZIP_EXPORT_GATES(tmp_path: Path):
    source = (ROOT / "desktop_host.py").read_text(encoding="utf-8")
    assert 'webview.settings["ALLOW_DOWNLOADS"] = True' in source
    assert source.index('webview.settings["ALLOW_DOWNLOADS"] = True') < source.index("webview.create_window")
    ui_source = (ROOT / "ui" / "rc1_app.py").read_text(encoding="utf-8")
    assert "导出 Word" in ui_source
    assert "导出单篇 ZIP" in ui_source
    assert "打开保存位置" in ui_source

    article = _method_article(_manual_topic("r1-2-export"))
    docx_path = export_article(article, tmp_path / "saved-word.docx")
    assert Document(docx_path).paragraphs[0].text == article["title"]
    zip_path = tmp_path / "saved-article.zip"
    export_article_bundle(article, tmp_path, zip_path)
    with zipfile.ZipFile(zip_path) as archive:
        names = archive.namelist()
        assert any(name.endswith(".docx") for name in names)
        assert any(name.endswith(".txt") for name in names)


def test_TEXT_KEY_LOAD_FAILED_WHEN_SAVED_KEY_FLAG_BUT_EMPTY_SECRET(tmp_path: Path, monkeypatch):
    topic = _manual_topic("r1-2-key-missing")
    monkeypatch.setattr(generation_store, "TASKS_ROOT", tmp_path / "tasks")
    store = SQLiteStore(tmp_path / "key.sqlite")
    task = _task(store, topic)
    result = single_task.run_single_task(
        task,
        {"api_key": "", "has_api_key": True, "model": "r1-2-text", "timeout_seconds": 30},
        {"api_key": "image-key", "model": "test-image"},
        settings={"network": {}, "image_plan_mode": "none"},
        store=store,
    )
    assert result["status"] == "failed"
    assert result["error_code"] == "TEXT_KEY_LOAD_FAILED"


def test_CUSTOM_TOPIC_SHORT_MODEL_OUTPUT_AUTO_EXPANDS_TO_FALLBACK_AND_WORD(tmp_path: Path, monkeypatch):
    """主链：手动话题 + 模型返回短文(<700字) → 自动扩写 → completed/warning + article非空 + gate≠failed + Word可导出"""
    topic = _manual_topic("r1-2-short-expand")
    seen = {}

    def fake_generate_short(topic, angle, article_type, style, word_count, profile, **kwargs):
        seen["called"] = True
        seen["api_key"] = profile.get("api_key")
        # Simulate model returning a very short article (well under 700 chars)
        return {
            "title": f"{topic.title}的简短思考",
            "intro": "AI赚钱是可行的。",
            "summary": topic.summary,
            "sections": [
                {"heading": "核心概念", "body": "用AI赚钱就是卖服务。", "image_brief": "概念"},
                {"heading": "方法", "body": "多尝试多学习。", "image_brief": "方法"},
            ],
            "content_markdown": f"# {topic.title}\n\nAI赚钱可行。\n\n## 核心概念\n用AI赚钱就是卖服务。\n\n## 方法\n多尝试多学习。",
            "source_list": [],
            "ai_statement": "AI辅助声明",
            "fact_basis": [],
            "recommended_status": "completed",
            "text_generation_calls": 1,
            "text_generation_limit": 1,
        }

    result = _run(tmp_path, monkeypatch, topic, fake_generate_short)
    article = result["article"]
    markdown = article["content_markdown"]

    # Core invariants: must not fail
    assert result["status"] in ("completed", "warning"), f"Expected completed/warning, got {result['status']}"
    assert result["article"] is not None
    assert result["quality_gate"]["status"] != "failed", f"Quality gate must not be failed: {result['quality_gate']}"

    # Auto-expand evidence
    assert result["fallback_kind"] == "custom_topic_expanded"
    assert "BODY_TOO_SHORT_EXPANDED" in str(result.get("fallback_reason", "")) or "BODY_TOO_SHORT_EXPANDED" in str(article.get("fallback_reason", ""))
    assert result["fallback_notice"]  # should have a notice about expansion

    # Method article structure
    assert "核心概念" in markdown
    assert "可执行方法" in markdown
    assert "具体步骤" in markdown
    assert "风险提醒" in markdown
    assert "总结" in markdown

    # Body length >= 700
    from generation.image_budget import count_body_chinese_chars
    body_chars = count_body_chinese_chars(article)
    assert body_chars >= 700, f"Expanded body must be >= 700, got {body_chars}"

    # Word export
    output = export_article(article, tmp_path / "short-expand.docx")
    assert Document(output).paragraphs[0].text == article["title"]


def _seed_export_state(tmp_path: Path, monkeypatch, *, task_status: str = "completed", article: dict | None = None, gate: dict | None = None):
    monkeypatch.setattr(generation_store, "TASKS_ROOT", tmp_path / "tasks")
    store = SQLiteStore(tmp_path / "export-gate.sqlite")
    topic = _manual_topic("r1-2-export-gate")
    task = _task(store, topic)
    store.update_task_status(task["task_id"], task_status)
    state = {
        "task_id": task["task_id"],
        "status": task_status,
        "stage": task_status,
        "state_version": 0,
        "article": article if article is not None else _method_article(topic),
        "quality_gate": gate if gate is not None else {"status": "passed", "passed": True, "hard_error_count": 0},
    }
    save_generation_task(state, allow_terminal_recovery=True)
    monkeypatch.setattr(api, "store", store)
    return task, state


def test_FAILED_TASK_CANNOT_EXPORT_EMPTY_WORD(tmp_path: Path, monkeypatch):
    task, _ = _seed_export_state(tmp_path, monkeypatch, task_status="failed")
    try:
        api._article_export(task["task_id"], "word")
    except ProviderError as exc:
        assert exc.code == "ARTICLE_NOT_READY"
    else:
        raise AssertionError("failed task must not export")
    assert not list((tmp_path / "tasks").rglob("*.docx"))


def test_EMPTY_BODY_MARKDOWN_CANNOT_EXPORT_WORD(tmp_path: Path, monkeypatch):
    article = {"title": "空正文测试", "intro": "导语存在", "sections": [], "content_markdown": "# 空正文测试"}
    task, _ = _seed_export_state(tmp_path, monkeypatch, article=article)
    try:
        api._article_export(task["task_id"], "word")
    except ProviderError as exc:
        assert exc.code == "ARTICLE_NOT_READY"
    else:
        raise AssertionError("empty article must not export")


def test_NOT_CHECKED_QUALITY_GATE_CANNOT_EXPORT_WORD(tmp_path: Path, monkeypatch):
    task, _ = _seed_export_state(tmp_path, monkeypatch, gate={"status": "not_checked", "hard_error_count": 0})
    try:
        api._article_export(task["task_id"], "word")
    except ProviderError as exc:
        assert exc.code == "ARTICLE_NOT_READY"
    else:
        raise AssertionError("not_checked gate must not export")


def test_FAILED_QUALITY_GATE_CANNOT_EXPORT_WORD(tmp_path: Path, monkeypatch):
    task, _ = _seed_export_state(tmp_path, monkeypatch, gate={"status": "failed", "hard_error_count": 1})
    try:
        api._article_export(task["task_id"], "word")
    except ProviderError as exc:
        assert exc.code == "ARTICLE_NOT_READY"
    else:
        raise AssertionError("failed gate must not export")


def test_PASSED_ARTICLE_CAN_EXPORT_WORD(tmp_path: Path, monkeypatch):
    task, _ = _seed_export_state(tmp_path, monkeypatch)
    response = api._article_export(task["task_id"], "word")
    assert Path(response.path).is_file()
    assert Document(response.path).paragraphs[0].text


def test_ZIP_DOES_NOT_INCLUDE_FAILED_EMPTY_WORD(tmp_path: Path, monkeypatch):
    article = {"title": "空正文测试", "intro": "导语存在", "sections": [], "content_markdown": "# 空正文测试"}
    task, _ = _seed_export_state(tmp_path, monkeypatch, task_status="failed", article=article, gate={"status": "not_checked"})
    try:
        api._article_export(task["task_id"], "zip")
    except ProviderError as exc:
        assert exc.code == "ARTICLE_NOT_READY"
    else:
        raise AssertionError("failed empty task must not export zip")
